"""
Research Conductor — Streamlit UI
Live pipeline: topic_planner -> phase_1 (parallel) -> rag_indexer -> critic_1 -> writer <> critic_2
"""

import asyncio
import base64
import concurrent.futures
import io
import json
import os
import tempfile
import time

import streamlit as st
from dotenv import load_dotenv

load_dotenv()

# Secrets bridge: copy st.secrets values into os.environ for agents.py
try:
    for _k in [
        "NVIDIA_API_KEY", "QDRANT_URL", "QDRANT_API_KEY", "QDRANT_COLLECTION",
        "SUPABASE_URL", "SUPABASE_ANON_KEY", "EMBEDDINGS_MODEL",
        "LLM_FAST", "LLM_WRITER", "LLM_CRITIC",
        "TAVILY_API_KEY", "REDIS_URL",
    ]:
        if _k in st.secrets and not os.getenv(_k):
            os.environ[_k] = st.secrets[_k]
except Exception:
    pass

# Auth gate — only active when [auth] is present in secrets.toml
_auth_configured = ("auth" in st.secrets) if hasattr(st, "secrets") else False
if _auth_configured:
    if not st.user.is_logged_in:
        st.set_page_config(page_title="Research Conductor — Login", layout="centered")
        st.title("Research Conductor")
        st.login()
        st.stop()
    _current_user_email: str = st.user.email or "anonymous"
else:
    _current_user_email: str = "anonymous"

import progress as _prog
from state import initial_state
from graph import build_graph

# ── Page config ────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Research Conductor",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ── Global CSS ─────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');

html, body, [data-testid="stAppViewContainer"] {
    font-family: 'Inter', -apple-system, sans-serif;
    background: #F8FAFC;
    color: #0F172A;
}
#MainMenu, footer, header,
[data-testid="stToolbar"],
[data-testid="stDecoration"] { display: none !important; }
.block-container { padding: 2rem 2.5rem 3rem !important; max-width: 1100px !important; margin: 0 auto; }

/* diagram wrapper — shimmer animation while running */
.dw {
    position: relative; overflow: hidden;
    border-radius: 14px; border: 1px solid #E2E8F0;
    background: #fff; box-shadow: 0 1px 4px rgba(0,0,0,.05);
    margin-bottom: 1.25rem;
}
.dw img { width: 100%; display: block; }
.dw.running::after {
    content: ''; position: absolute; inset: 0;
    background: linear-gradient(90deg, transparent 20%, rgba(255,255,255,.55) 50%, transparent 80%);
    animation: shimmer 1.6s ease-in-out infinite;
    pointer-events: none;
}
@keyframes shimmer { 0%{transform:translateX(-100%)} 100%{transform:translateX(100%)} }

/* metric chips */
.mrow { display: flex; gap: 10px; flex-wrap: wrap; margin-bottom: 1.1rem; }
.mchip {
    background: #fff; border: 1px solid #E2E8F0;
    border-radius: 10px; padding: 10px 18px; min-width: 105px;
    box-shadow: 0 1px 2px rgba(0,0,0,.03);
}
.mlabel { font-size: 10px; font-weight: 700; letter-spacing: .08em;
          color: #94A3B8; text-transform: uppercase; }
.mval { font-size: 1.4rem; font-weight: 700; color: #0F172A;
        letter-spacing: -.015em; margin-top: 3px; }
.mval.g { color: #10B981; } .mval.o { color: #F59E0B; }
.mval.r { color: #EF4444; } .mval.p { color: #6366F1; }

/* result verdict badges */
.vbadge { display: inline-block; padding: 4px 14px;
          border-radius: 6px; font-weight: 600; font-size: 13px; }
.v-pass     { background:#D1FAE5; color:#065F46; }
.v-revise   { background:#FEF3C7; color:#92400E; }
.v-human    { background:#FEE2E2; color:#991B1B; }
.v-emerging { background:#EDE9FE; color:#4C1D95; }
.v-stable   { background:#D1FAE5; color:#065F46; }
.v-declining{ background:#FEF3C7; color:#92400E; }
.v-dead     { background:#FEE2E2; color:#991B1B; }

.issue-item { background:#FFF7ED; border-left:3px solid #F97316;
              padding:8px 12px; margin:4px 0; border-radius:4px;
              font-size:.875rem; }

.sec { font-size:10px; font-weight:700; letter-spacing:.1em;
       color:#94A3B8; text-transform:uppercase; margin:1.5rem 0 .6rem; }

/* active agents list */
.arow { display:flex; align-items:center; gap:8px; padding:4px 0;
        font-size:12px; color:#374151; }
.adot { width:8px; height:8px; border-radius:50%; flex-shrink:0; }
.atag { margin-left:auto; font-size:10px; font-weight:700; }

@media (prefers-reduced-motion: reduce) { .dw.running::after { animation:none; } }
</style>
""", unsafe_allow_html=True)

# ── Cached graph ───────────────────────────────────────────────────────────────
@st.cache_resource(show_spinner="Compiling pipeline...")
def get_graph():
    return build_graph()

# ── Persistent executor (survives reruns) ──────────────────────────────────────
if "executor" not in st.session_state:
    st.session_state.executor = concurrent.futures.ThreadPoolExecutor(max_workers=1)
for k, v in [("pipeline_log", []), ("pipeline_running", False), ("chat_history", [])]:
    if k not in st.session_state:
        st.session_state[k] = v


def _fresh_executor():
    """Shut down any stuck old executor and return a new one."""
    old = st.session_state.get("executor")
    if old is not None:
        old.shutdown(wait=False, cancel_futures=True)
    ex = concurrent.futures.ThreadPoolExecutor(max_workers=1)
    st.session_state.executor = ex
    return ex

# ── Agent colour table ─────────────────────────────────────────────────────────
COLORS = {
    "topic_planner": "#6366F1",
    "ingestion":     "#3B82F6",
    "currency":      "#F59E0B",
    "memory":        "#10B981",
    "rag":           "#06B6D4",
    "error_handler": "#EF4444",
    "critic_1":      "#F97316",
    "writer":        "#6366F1",
    "critic_2":      "#F97316",
}

try:
    from agents import MAX_RETRIES, MAX_ROUNDS
except Exception:
    MAX_RETRIES, MAX_ROUNDS = 2, 3


# ── Derive agent states from event log ────────────────────────────────────────
def agent_states(log: list) -> dict:
    s: dict[str, str] = {}
    for ev in log:
        ag, st2 = ev["agent"], ev["status"]
        if st2 == "start":
            s[ag] = "running"
        elif st2 == "done":
            s[ag] = "done"
        elif st2 == "error":
            s[ag] = "error"
        elif st2 == "warn" and ag not in s:
            s[ag] = "warn"
    return s


# ── SVG diagram (base64 image — bypasses Streamlit text-node stripping) ────────
def _node_colors(agent: str, states: dict):
    c = COLORS.get(agent, "#6B7280")
    s = states.get(agent, "idle")
    if s == "idle":    return "#F8FAFC", "#E2E8F0", "#94A3B8"
    if s == "running": return f"{c}22", c, c
    if s == "done":    return "#F0FDF4", "#86EFAC", "#15803D"
    if s == "error":   return "#FEF2F2", "#FCA5A5", "#B91C1C"
    if s == "warn":    return "#FFFBEB", "#FCD34D", "#92400E"
    return "#F8FAFC", "#E2E8F0", "#94A3B8"


def _n(x, y, w, h, agent, label, sub, states):
    fill, stroke, txt = _node_colors(agent, states)
    sw = "2" if states.get(agent) == "running" else "1.5"
    sub_el = (
        f'<text x="{x+w//2}" y="{y+h-5}" text-anchor="middle" '
        f'font-size="9" fill="{txt}" opacity=".75">{sub}</text>'
    ) if sub else ""
    return (
        f'<rect x="{x}" y="{y}" width="{w}" height="{h}" rx="8" '
        f'fill="{fill}" stroke="{stroke}" stroke-width="{sw}"/>'
        f'<text x="{x+w//2}" y="{y+h//2+5}" text-anchor="middle" '
        f'font-size="11" font-weight="600" fill="{txt}">{label}</text>'
        f'{sub_el}'
    )


def build_svg(states: dict, retry_count: int = 0,
              is_running: bool = False, is_done: bool = False) -> str:
    def sub(ag):
        s = states.get(ag)
        return "done" if s == "done" else ("..." if s == "running" else "")

    ac  = "#6366F1" if is_running else "#CBD5E1"
    dot = "#10B981" if is_done else ("#6366F1" if is_running else "#CBD5E1")

    nodes = (
        _n(28,  82, 90, 34, "topic_planner", "Planner",  sub("topic_planner"), states) +
        _n(140, 14, 90, 34, "ingestion",     "Ingestion", sub("ingestion"),    states) +
        _n(140, 82, 90, 34, "currency",      "Currency",  sub("currency"),     states) +
        _n(140,150, 90, 34, "memory",        "Memory",    sub("memory"),       states) +
        _n(258, 82, 80, 34, "rag",           "RAG",       sub("rag"),          states) +
        _n(352, 82, 80, 34, "critic_1",      "Critic 1",  sub("critic_1"),     states) +
        _n(446, 82, 80, 34, "writer",        "Writer",    sub("writer"),       states) +
        _n(540, 82, 80, 34, "critic_2",      "Critic 2",  sub("critic_2"),     states)
    )

    err_node = ""
    if retry_count > 0:
        fill, stroke, txt = _node_colors("error_handler", states)
        err_node = (
            f'<rect x="140" y="178" width="90" height="26" rx="6" '
            f'fill="{fill}" stroke="{stroke}" stroke-width="1.5" stroke-dasharray="4,2"/>'
            f'<text x="185" y="195" text-anchor="middle" font-size="10" '
            f'font-weight="600" fill="{txt}">Retry {retry_count}/{MAX_RETRIES}</text>'
        )

    done_badge = (
        '<text x="660" y="14" text-anchor="middle" font-size="9" '
        'font-weight="700" fill="#10B981" letter-spacing=".04em">DONE</text>'
    ) if is_done else ""

    # Arrow marker color in SVG defs must be inline — use a unique id per color
    m_id = "arr"
    loop_id = "larr"

    return f"""<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 700 215"
     style="font-family:Inter,-apple-system,sans-serif">
  <defs>
    <marker id="{m_id}" markerWidth="7" markerHeight="6"
            refX="7" refY="3" orient="auto">
      <polygon points="0 0,7 3,0 6" fill="{ac}"/>
    </marker>
    <marker id="{loop_id}" markerWidth="7" markerHeight="6"
            refX="7" refY="3" orient="auto">
      <polygon points="0 0,7 3,0 6" fill="#F97316"/>
    </marker>
  </defs>

  <!-- start dot -->
  <circle cx="14" cy="99" r="8" fill="{dot}"/>
  <circle cx="14" cy="99" r="3" fill="white"/>

  <!-- start to planner -->
  <line x1="22" y1="99" x2="27" y2="99" stroke="{ac}" stroke-width="1.5"
        marker-end="url(#{m_id})"/>

  <!-- planner to fork -->
  <line x1="118" y1="99" x2="128" y2="99" stroke="{ac}" stroke-width="1.5"/>
  <line x1="128" y1="31"  x2="128" y2="167" stroke="{ac}" stroke-width="1.5"/>
  <line x1="128" y1="31"  x2="139" y2="31"  stroke="{ac}" stroke-width="1.5" marker-end="url(#{m_id})"/>
  <line x1="128" y1="99"  x2="139" y2="99"  stroke="{ac}" stroke-width="1.5" marker-end="url(#{m_id})"/>
  <line x1="128" y1="167" x2="139" y2="167" stroke="{ac}" stroke-width="1.5" marker-end="url(#{m_id})"/>

  <!-- phase 1 to join -->
  <line x1="230" y1="31"  x2="243" y2="31"  stroke="{ac}" stroke-width="1.5"/>
  <line x1="230" y1="99"  x2="243" y2="99"  stroke="{ac}" stroke-width="1.5"/>
  <line x1="230" y1="167" x2="243" y2="167" stroke="{ac}" stroke-width="1.5"/>
  <line x1="243" y1="31"  x2="243" y2="167" stroke="{ac}" stroke-width="1.5"/>
  <line x1="243" y1="99"  x2="257" y2="99"  stroke="{ac}" stroke-width="1.5" marker-end="url(#{m_id})"/>

  <!-- rag to critic1 to writer to critic2 to end -->
  <line x1="338" y1="99" x2="351" y2="99" stroke="{ac}" stroke-width="1.5" marker-end="url(#{m_id})"/>
  <line x1="432" y1="99" x2="445" y2="99" stroke="{ac}" stroke-width="1.5" marker-end="url(#{m_id})"/>
  <line x1="526" y1="99" x2="539" y2="99" stroke="{ac}" stroke-width="1.5" marker-end="url(#{m_id})"/>
  <line x1="620" y1="99" x2="636" y2="99" stroke="{ac}" stroke-width="1.5" marker-end="url(#{m_id})"/>

  <!-- revise loop -->
  <path d="M580,116 V158 H486 V116" fill="none"
        stroke="#F97316" stroke-width="1.5" stroke-dasharray="4,3"
        marker-end="url(#{loop_id})"/>
  <text x="533" y="173" text-anchor="middle" font-size="9"
        fill="#F97316" font-style="italic">revise</text>

  <!-- end dot -->
  <circle cx="648" cy="99" r="8" fill="{dot}"/>
  <circle cx="648" cy="99" r="4" fill="white"/>
  <circle cx="648" cy="99" r="2" fill="{dot}"/>

  <!-- column labels -->
  <text x="73"  y="8" text-anchor="middle" font-size="8" fill="#CBD5E1"
        font-weight="700" letter-spacing=".07em">PLANNER</text>
  <text x="185" y="8" text-anchor="middle" font-size="8" fill="#CBD5E1"
        font-weight="700" letter-spacing=".07em">PHASE 1</text>
  <text x="298" y="8" text-anchor="middle" font-size="8" fill="#CBD5E1"
        font-weight="700" letter-spacing=".07em">RAG</text>
  <text x="392" y="8" text-anchor="middle" font-size="8" fill="#CBD5E1"
        font-weight="700" letter-spacing=".07em">CRITIC 1</text>
  <text x="486" y="8" text-anchor="middle" font-size="8" fill="#CBD5E1"
        font-weight="700" letter-spacing=".07em">WRITER</text>
  <text x="580" y="8" text-anchor="middle" font-size="8" fill="#CBD5E1"
        font-weight="700" letter-spacing=".07em">CRITIC 2</text>

  {nodes}
  {err_node}
  {done_badge}
</svg>"""


@st.cache_data(show_spinner=False)
def _build_svg_cached(states_key: tuple, retry_count: int,
                      is_running: bool, is_done: bool) -> str:
    states = dict(states_key)
    return build_svg(states, retry_count, is_running, is_done)


def render_diagram(states: dict, retry_count: int = 0,
                   is_running: bool = False, is_done: bool = False) -> str:
    """SVG embedded as base64 image — immune to Streamlit's HTML sanitizer."""
    states_key = tuple(sorted(states.items()))
    svg = _build_svg_cached(states_key, retry_count, is_running, is_done)
    b64 = base64.b64encode(svg.encode("utf-8")).decode()
    cls = "dw running" if is_running else "dw"
    return (
        f'<div class="{cls}">'
        f'<img src="data:image/svg+xml;base64,{b64}" '
        f'style="width:100%;display:block;padding:10px 14px;box-sizing:border-box">'
        f'</div>'
    )


# ── Live log (rendered via st.components.v1.html — iframe, full fidelity) ─────
def render_live_log_html(log: list, running: bool = True) -> str:
    """Returns a full self-contained HTML page for st.components.v1.html()."""
    BADGE = {
        "topic_planner": ("8B5CF6", "EDE9FE"),
        "ingestion":     ("3B82F6", "DBEAFE"),
        "currency":      ("D97706", "FEF3C7"),
        "memory":        ("059669", "D1FAE5"),
        "rag":           ("0891B2", "CFFAFE"),
        "error_handler": ("DC2626", "FEE2E2"),
        "critic_1":      ("EA580C", "FFF7ED"),
        "writer":        ("4F46E5", "EDE9FE"),
        "critic_2":      ("EA580C", "FFF7ED"),
    }
    SMSG = {"done": "4ADE80", "warn": "FB923C", "error": "F87171"}
    PREFIX = {"start": "&gt;&gt;", "done": "OK", "info": "&nbsp;&nbsp;", "warn": "**", "error": "!!"}

    lines = ""
    for ev in log[-80:]:
        ag  = ev.get("agent", "system")
        st2 = ev.get("status", "info")
        msg = ev.get("msg", "").replace("<", "&lt;").replace(">", "&gt;")
        ts  = ev.get("ts", "")
        fc, bc = BADGE.get(ag, ("64748B", "F1F5F9"))
        label  = _prog.AGENT_LABELS.get(ag, ag)
        pfx    = PREFIX.get(st2, "&nbsp;&nbsp;")
        mcol   = SMSG.get(st2, "94A3B8")
        lines += (
            f'<div class="ll">'
            f'<span class="ts">{ts}</span>'
            f'<span class="badge" style="background:#{bc};color:#{fc}">{label}</span>'
            f'<span style="color:#{mcol}">{pfx}&nbsp;{msg}</span>'
            f'</div>\n'
        )

    cursor = '<span class="cur"></span>' if running else ""
    status_text = "running..." if running else "pipeline complete"

    return f"""<!DOCTYPE html><html><head><meta charset="utf-8">
<style>
* {{ box-sizing:border-box; margin:0; padding:0; }}
body {{ background:#0F172A; font-family:'JetBrains Mono',Menlo,monospace;
        font-size:12px; color:#CBD5E1; overflow:hidden; }}
.tbar {{ background:#1E293B; padding:9px 14px; display:flex; align-items:center;
         gap:7px; border-bottom:1px solid #334155; }}
.d  {{ width:10px;height:10px;border-radius:50%;display:inline-block; }}
.dr {{ background:#EF4444; }} .dy {{ background:#F59E0B; }} .dg {{ background:#10B981; }}
.tname {{ font-size:11px; color:#64748B; margin-left:3px; }}
.body  {{ padding:12px 16px; height:calc(100vh - 40px);
          overflow-y:auto; line-height:1.75; }}
.ll    {{ display:flex; gap:10px; align-items:baseline;
          animation:lin 150ms ease both; }}
@keyframes lin {{ from{{opacity:0;transform:translateX(-3px)}} to{{opacity:1;transform:translateX(0)}} }}
.ts    {{ color:#334155; flex-shrink:0; font-size:10.5px; }}
.badge {{ flex-shrink:0; padding:1px 8px; border-radius:4px;
          font-size:10px; font-weight:700; letter-spacing:.04em; }}
.cur   {{ display:inline-block; width:7px; height:13px;
          background:#6366F1; vertical-align:middle; margin-left:3px;
          animation:blink 1s step-end infinite; }}
@keyframes blink {{ 50%{{opacity:0}} }}
@media (prefers-reduced-motion:reduce) {{ .ll{{animation:none}} .cur{{animation:none}} }}
</style></head>
<body>
<div class="tbar">
  <span class="d dr"></span><span class="d dy"></span><span class="d dg"></span>
  <span class="tname">pipeline activity</span>
</div>
<div class="body" id="b">
{lines}
<span style="color:#334155">{status_text}</span>{cursor}
</div>
<script>var b=document.getElementById('b');b.scrollTop=b.scrollHeight;</script>
</body></html>"""


# ── Metric chips ───────────────────────────────────────────────────────────────
def _latest(log: list, key: str, default=None):
    for ev in reversed(log):
        if key in ev:
            return ev[key]
    return default


def render_metrics(log: list) -> str:
    papers  = _latest(log, "count", "-")
    verdict = _latest(log, "verdict", "-")
    q_num   = _latest(log, "query_num", 0)
    q_tot   = _latest(log, "query_total", 0)
    retry   = _latest(log, "retry", 0)
    rnd     = _latest(log, "round", 0)

    v_cls = {"EMERGING": "p", "STABLE": "g", "DECLINING": "o", "DEAD": "r"}.get(str(verdict), "")

    def chip(label, val, cls=""):
        return (f'<div class="mchip"><div class="mlabel">{label}</div>'
                f'<div class="mval {cls}">{val}</div></div>')

    return (
        f'<div class="mrow">'
        f'{chip("Papers", papers, "p")}'
        f'{chip("Queries", f"{q_num}/{q_tot}" if q_tot else "-")}'
        f'{chip("Retries", f"{retry}/{MAX_RETRIES}", "o" if retry > 0 else "")}'
        f'{chip("Round", f"{rnd}/{MAX_ROUNDS}")}'
        f'{chip("Currency", verdict, v_cls)}'
        f'</div>'
    )


# ── Export helpers ─────────────────────────────────────────────────────────────
def _export_json(result: dict) -> bytes:
    payload = {
        "topic":    result.get("topic"),
        "plan":     result.get("search_plan"),
        "verdict":  result.get("final_verdict"),
        "confidence": result.get("confidence"),
        "draft":    result.get("draft"),
        "feedback": result.get("critic_feedback"),
        "papers":   result.get("papers"),
    }
    return json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")


def _export_pdf(result: dict) -> bytes:
    try:
        from fpdf import FPDF
    except ImportError:
        return b""
    topic = result.get("topic", "Research")
    draft = result.get("draft", "")
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.multi_cell(0, 10, topic)
    pdf.ln(4)
    pdf.set_font("Helvetica", "", 11)
    for para in draft.split("\n\n"):
        cleaned = para.strip()
        if cleaned:
            pdf.multi_cell(0, 7, cleaned)
            pdf.ln(3)
    return bytes(pdf.output())


def _export_docx(result: dict) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        return b""
    topic = result.get("topic", "Research")
    draft = result.get("draft", "")
    papers = result.get("papers", [])
    doc = Document()
    doc.add_heading(topic, level=1)
    for para in draft.split("\n\n"):
        cleaned = para.strip()
        if cleaned:
            doc.add_paragraph(cleaned)
    if papers:
        doc.add_heading("References", level=2)
        for p in papers:
            doc.add_paragraph(
                f"{p.get('title','Untitled')} ({p.get('year','?')}) — {p.get('source','')}",
                style="List Bullet",
            )
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Results display ────────────────────────────────────────────────────────────
def render_results(result: dict):
    plan   = result.get("search_plan") or {}
    papers = result.get("papers", [])
    srch   = result.get("search_results", [])
    verd   = result.get("currency_verdict", "UNKNOWN")
    score  = result.get("currency_score", 0.0)
    reason = result.get("currency_reason", "")
    memory = result.get("memory_context", "")
    notes  = result.get("critic1_notes", "")
    fv     = result.get("final_verdict", "UNKNOWN")
    rounds = result.get("round_num", 0)
    conf   = result.get("confidence", 0.0)
    issues = result.get("critic_feedback", [])
    draft  = result.get("draft", "")
    p1err  = result.get("phase1_error")
    retry  = result.get("retry_count", 0)
    topic  = result.get("topic", "")

    # Topic plan
    if plan:
        st.markdown('<div class="sec">Topic Plan</div>', unsafe_allow_html=True)
        c1, c2 = st.columns(2)
        with c1:
            st.markdown(f"**Year range:** {plan.get('year_from','?')} - {plan.get('year_to','?')}")
            if plan.get("rationale"):
                st.caption(plan["rationale"])
        with c2:
            if plan.get("aspects"):
                st.markdown("**Aspects:** " + " / ".join(plan["aspects"]))
        if plan.get("queries"):
            with st.expander(f"{len(plan['queries'])} search sub-queries"):
                for i, q in enumerate(plan["queries"], 1):
                    st.markdown(f"`{i}.` {q}")

    if p1err:
        st.warning(f"Orchestrator recovery ({retry}x retries): {p1err}")

    # Phase 1
    st.markdown('<div class="sec">Phase 1 - Ingestion / Currency / Memory</div>', unsafe_allow_html=True)
    m1, m2, m3, m4 = st.columns(4)
    m1.metric("Papers ingested", len(papers))
    m2.metric("Search results", len(srch))
    v_cls = {"EMERGING":"v-emerging","STABLE":"v-stable",
             "DECLINING":"v-declining","DEAD":"v-dead"}.get(verd, "v-revise")
    m3.markdown(
        f'<div style="padding-top:6px"><div style="font-size:10px;font-weight:700;'
        f'letter-spacing:.08em;color:#94A3B8;text-transform:uppercase;margin-bottom:6px">'
        f'Currency</div><span class="vbadge {v_cls}">{verd}</span></div>',
        unsafe_allow_html=True
    )
    m4.metric("Score", f"{score:.2f}")
    if reason: st.caption(f"Reason: {reason}")
    if memory: st.caption(f"Memory: {memory}")
    if papers:
        with st.expander(f"Papers ingested ({len(papers)})"):
            for p in papers:
                cites = f" / {p.get('citations','')} cites" if p.get("citations") else ""
                st.markdown(
                    f"- **{p.get('title','Untitled')}** ({p.get('year','?')}) "
                    f"_{p.get('source','')}{cites}_"
                )

    # Critic 1
    st.markdown('<div class="sec">Critic 1</div>', unsafe_allow_html=True)
    ok = "reasonable" in notes.lower()
    (st.success if ok else st.warning)(notes or "No notes.")

    # Debate
    st.markdown('<div class="sec">Writer / Critic 2 Debate</div>', unsafe_allow_html=True)
    d1, d2, d3 = st.columns(3)
    d1.metric("Rounds", rounds)
    d2.metric("Confidence", f"{(conf or 0):.0%}")
    fv_cls = {"PASS":"v-pass","HUMAN_REVIEW":"v-human"}.get(fv, "v-revise")
    d3.markdown(
        f'<div style="padding-top:6px"><span class="vbadge {fv_cls}">{fv}</span></div>',
        unsafe_allow_html=True
    )
    if issues:
        with st.expander(f"Outstanding issues ({len(issues)}) — click to fix in chat",
                         expanded=(fv == "HUMAN_REVIEW")):
            # "Fix all" header button
            all_issues_prompt = (
                "The research draft has the following outstanding issues from the academic reviewer. "
                "Please rewrite the relevant sections of the draft to address ALL of them:\n\n"
                + "\n".join(f"{i+1}. {iss}" for i, iss in enumerate(issues))
            )
            if st.button("Fix all issues in chat", type="primary", key="fix_all_issues"):
                st.session_state.chat_history.append({"role": "user", "content": all_issues_prompt})
                st.rerun()

            st.markdown("<div style='height:6px'></div>", unsafe_allow_html=True)

            # Individual issues with per-issue fix button
            for i, iss in enumerate(issues):
                ic1, ic2 = st.columns([6, 1])
                ic1.markdown(f'<div class="issue-item">{iss}</div>', unsafe_allow_html=True)
                fix_prompt = (
                    f"The academic reviewer flagged this specific issue with the draft:\n\n"
                    f"\"{iss}\"\n\n"
                    "Please rewrite or expand the relevant part of the research draft to fix this issue, "
                    "citing from the available papers."
                )
                if ic2.button("Fix", key=f"fix_issue_{i}", use_container_width=True):
                    st.session_state.chat_history.append({"role": "user", "content": fix_prompt})
                    st.rerun()

    # Draft
    st.markdown('<div class="sec">Research Snapshot</div>', unsafe_allow_html=True)
    if draft:
        st.markdown(draft)
        slug = topic[:35].replace(" ", "_")
        ec1, ec2, ec3, ec4 = st.columns(4)
        ec1.download_button(
            "Download TXT", data=draft,
            file_name=f"{slug}.txt", mime="text/plain",
            use_container_width=True,
        )
        ec2.download_button(
            "Download JSON", data=_export_json(result),
            file_name=f"{slug}.json", mime="application/json",
            use_container_width=True,
        )
        pdf_bytes = _export_pdf(result)
        if pdf_bytes:
            ec3.download_button(
                "Download PDF", data=pdf_bytes,
                file_name=f"{slug}.pdf", mime="application/pdf",
                use_container_width=True,
            )
        else:
            ec3.caption("PDF unavailable")
        docx_bytes = _export_docx(result)
        if docx_bytes:
            ec4.download_button(
                "Download Word", data=docx_bytes,
                file_name=f"{slug}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        else:
            ec4.caption("Word unavailable")
    else:
        st.warning("No draft produced.")


# ── Pipeline runner (background thread) ───────────────────────────────────────
def _run(state_dict: dict) -> dict:
    loop = asyncio.new_event_loop()
    asyncio.set_event_loop(loop)
    try:
        return loop.run_until_complete(get_graph().ainvoke(state_dict))
    finally:
        loop.close()


# ══════════════════════════════════════════════════════════════════════════════
# APP LAYOUT
# ══════════════════════════════════════════════════════════════════════════════

st.markdown(
    '<h1 style="font-size:1.7rem;font-weight:700;letter-spacing:-.02em;margin-bottom:2px">'
    'Research Conductor</h1>'
    '<p style="font-size:.875rem;color:#64748B;margin-bottom:1.4rem">'
    'Agentic pipeline &nbsp;|&nbsp; Multi-source ingestion &nbsp;|&nbsp; '
    'RAG grounding &nbsp;|&nbsp; Writer / Critic debate</p>',
    unsafe_allow_html=True,
)

# ── Sidebar (auth info + logout) ───────────────────────────────────────────────
if _auth_configured:
    with st.sidebar:
        st.markdown(f"**Signed in as**  \n{_current_user_email}")
        st.logout()

# ── Inputs ─────────────────────────────────────────────────────────────────────
ic, bc = st.columns([5, 1])
with ic:
    topic = st.text_input("topic", placeholder="e.g. Retrieval-Augmented Generation for medical literature",
                          label_visibility="collapsed")
with bc:
    st.markdown("<div style='height:6px'></div>", unsafe_allow_html=True)
    run_btn = st.button("Run Research", type="primary", use_container_width=True,
                        disabled=not topic.strip() or st.session_state.pipeline_running)

uploaded = st.file_uploader("PDFs (optional)", type=["pdf"],
                             accept_multiple_files=True, label_visibility="visible")

st.markdown("---", unsafe_allow_html=False)

# ── Derive current state ────────────────────────────────────────────────────────
is_running = st.session_state.pipeline_running
has_result = "pipeline_result" in st.session_state
log        = st.session_state.pipeline_log
states     = agent_states(log)
retry_cnt  = _latest(log, "retry", 0)

# ── Diagram (always visible, updates each rerun) ───────────────────────────────
st.markdown(
    render_diagram(states, retry_count=retry_cnt,
                   is_running=is_running, is_done=has_result),
    unsafe_allow_html=True,
)

# ── LIVE UI (while running) ────────────────────────────────────────────────────
if is_running and "pipeline_future" in st.session_state:
    future = st.session_state.pipeline_future

    # Drain new events
    new_evs = _prog.drain()
    st.session_state.pipeline_log.extend(new_evs)
    log    = st.session_state.pipeline_log
    states = agent_states(log)

    # Live metrics
    st.markdown(render_metrics(log), unsafe_allow_html=True)

    # Log + active agents
    lc, rc = st.columns([3, 1])
    with lc:
        st.components.v1.html(render_live_log_html(log, running=True), height=370, scrolling=False)
    with rc:
        st.markdown("**Active agents**")
        for ag, ag_st in states.items():
            c = COLORS.get(ag, "#6B7280")
            lbl = _prog.AGENT_LABELS.get(ag, ag)
            icon = {"running": "...", "done": "OK", "error": "!!", "warn": "**"}.get(ag_st, "")
            st.markdown(
                f'<div class="arow">'
                f'<span class="adot" style="background:{c}"></span>'
                f'{lbl}'
                f'<span class="atag" style="color:{c}">{icon}</span>'
                f'</div>',
                unsafe_allow_html=True,
            )

    # Check if done
    if future.done():
        for p in st.session_state.get("pipeline_paper_paths", []):
            try: os.unlink(p)
            except Exception: pass
        try:
            r = future.result()
            r["topic"] = st.session_state.get("pipeline_topic", "")
            st.session_state.pipeline_result = r
        except Exception as e:
            st.session_state.pipeline_error = str(e)
        del st.session_state.pipeline_future
        st.session_state.pipeline_running = False
        st.rerun()
    else:
        time.sleep(0.8)
        st.rerun()

# ── ERROR ──────────────────────────────────────────────────────────────────────
elif "pipeline_error" in st.session_state:
    st.error(f"Pipeline error: {st.session_state.pipeline_error}")
    if st.button("Clear"):
        del st.session_state.pipeline_error
        st.rerun()

# ── RESULTS ────────────────────────────────────────────────────────────────────
elif has_result:
    with st.expander("Activity log", expanded=False):
        st.components.v1.html(render_live_log_html(log, running=False), height=370, scrolling=False)

    st.divider()
    render_results(st.session_state.pipeline_result)
    st.divider()

    # ── THESIS CHAT ───────────────────────────────────────────────────────────
    st.markdown("""
<style>
/* ── Chat container ── */
.chat-wrap { border:1px solid #E2E8F0; border-radius:14px; overflow:hidden;
             background:#fff; box-shadow:0 1px 4px rgba(0,0,0,.05); margin-top:1.5rem; }
.chat-header { padding:14px 20px; border-bottom:1px solid #F1F5F9;
               display:flex; align-items:center; justify-content:space-between; }
.chat-title  { font-size:.95rem; font-weight:700; color:#0F172A;
               letter-spacing:-.01em; }
.chat-sub    { font-size:.75rem; color:#94A3B8; margin-top:1px; }
.chat-body   { padding:20px; display:flex; flex-direction:column; gap:14px; }

/* message bubbles */
.msg         { display:flex; gap:10px; align-items:flex-start; max-width:100%; }
.msg.user    { flex-direction:row-reverse; }
.avatar      { width:30px; height:30px; border-radius:50%; flex-shrink:0;
               display:flex; align-items:center; justify-content:center;
               font-size:12px; font-weight:700; letter-spacing:0; }
.av-user     { background:#EDE9FE; color:#6366F1; }
.av-ai       { background:#F0FDF4; color:#059669; border:1px solid #BBF7D0; }
.bubble      { padding:11px 15px; border-radius:12px; font-size:.85rem;
               line-height:1.65; max-width:88%; white-space:pre-wrap; }
.bubble-user { background:#6366F1; color:#fff; border-bottom-right-radius:4px; }
.bubble-ai   { background:#F8FAFC; color:#1E293B; border:1px solid #E2E8F0;
               border-bottom-left-radius:4px; }

/* quick-start pills */
.qrow  { display:flex; flex-wrap:wrap; gap:8px; padding:14px 20px;
          border-top:1px solid #F1F5F9; }
.qpill { font-size:.75rem; font-weight:600; color:#6366F1;
         background:#EDE9FE; border:none; border-radius:20px;
         padding:5px 13px; cursor:pointer; transition:background .15s; }
.qpill:hover { background:#DDD6FE; }
</style>
""", unsafe_allow_html=True)

    QUICK = [
        ("Literature Review",  "Write a detailed literature review section for my thesis on this topic, citing the available papers."),
        ("Methodology",        "Suggest a methodology section for a thesis on this topic based on approaches used in the retrieved papers."),
        ("Introduction",       "Write an academic introduction section for a thesis on this topic, ending with a clear research gap statement."),
        ("Discussion",         "Write a discussion section that interprets the research snapshot findings and situates them in the broader literature."),
        ("Conclusion",         "Write a conclusion section that summarises key findings, limitations, and directions for future research."),
        ("Bibliography",       "Generate a bibliography / reference list from the available papers in APA 7th edition format."),
    ]

    # ── Chat container header ─────────────────────────────────────────────────
    st.markdown(
        '<div class="chat-wrap">'
        '<div class="chat-header">'
        '<div><div class="chat-title">Thesis Writing Assistant</div>'
        '<div class="chat-sub">Grounded in your research papers and snapshot</div></div>'
        '</div>',
        unsafe_allow_html=True,
    )

    # ── Render existing history ───────────────────────────────────────────────
    if st.session_state.chat_history:
        history_all = st.session_state.chat_history
        PAGE_SIZE = 20
        if len(history_all) > PAGE_SIZE:
            with st.expander(f"Show {len(history_all) - PAGE_SIZE} earlier messages"):
                older_html = '<div class="chat-body">'
                for msg in history_all[:-PAGE_SIZE]:
                    content = msg["content"].replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    role_cls = "user" if msg["role"] == "user" else ""
                    av_cls   = "av-user" if msg["role"] == "user" else "av-ai"
                    av_lbl   = "You" if msg["role"] == "user" else "AI"
                    bubble   = "bubble-user" if msg["role"] == "user" else "bubble-ai"
                    older_html += (
                        f'<div class="msg {role_cls}">'
                        f'<span class="avatar {av_cls}">{av_lbl}</span>'
                        f'<div class="bubble {bubble}">{content}</div>'
                        f'</div>'
                    )
                older_html += '</div>'
                st.markdown(older_html, unsafe_allow_html=True)
        msgs_html = '<div class="chat-body">'
        for msg in history_all[-PAGE_SIZE:]:
            content = msg["content"].replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            if msg["role"] == "user":
                msgs_html += (
                    f'<div class="msg user">'
                    f'<span class="avatar av-user">You</span>'
                    f'<div class="bubble bubble-user">{content}</div>'
                    f'</div>'
                )
            else:
                msgs_html += (
                    f'<div class="msg">'
                    f'<span class="avatar av-ai">AI</span>'
                    f'<div class="bubble bubble-ai">{content}</div>'
                    f'</div>'
                )
        msgs_html += '</div>'
        st.markdown(msgs_html, unsafe_allow_html=True)

    # ── Quick-start pills (inside card) ──────────────────────────────────────
    st.markdown('<div class="qrow">', unsafe_allow_html=True)
    qs_cols = st.columns(6)
    for i, (label, prompt) in enumerate(QUICK):
        if qs_cols[i].button(label, key=f"qs_{i}", use_container_width=True):
            st.session_state.chat_history.append({"role": "user", "content": prompt})
            st.rerun()
    st.markdown('</div></div>', unsafe_allow_html=True)  # close qrow + chat-wrap

    # ── Chat input ────────────────────────────────────────────────────────────
    user_input = st.chat_input("Ask for a thesis section, outline, or refinement...")

    def _get_response(user_msg: str) -> str:
        from agents import chat_with_research
        return chat_with_research(
            user_message=user_msg,
            result=st.session_state.pipeline_result,
            history=st.session_state.chat_history[:-1],
        )

    if user_input:
        st.session_state.chat_history.append({"role": "user", "content": user_input})
        with st.spinner("Writing..."):
            response = _get_response(user_input)
        st.session_state.chat_history.append({"role": "assistant", "content": response})
        st.rerun()

    # Pending quick-start (button press triggers rerun before LLM call)
    last = st.session_state.chat_history
    if last and last[-1]["role"] == "user" and not user_input:
        with st.spinner("Writing..."):
            response = _get_response(last[-1]["content"])
        st.session_state.chat_history.append({"role": "assistant", "content": response})
        st.rerun()

    # ── Footer controls ───────────────────────────────────────────────────────
    st.markdown("<div style='height:10px'></div>", unsafe_allow_html=True)
    cc, nc = st.columns([1, 1])
    if cc.button("Clear chat", type="secondary"):
        st.session_state.chat_history = []
        st.rerun()
    if nc.button("New research", type="secondary"):
        for k in ["pipeline_result", "pipeline_log", "pipeline_error", "chat_history"]:
            st.session_state.pop(k, None)
        _prog.clear()
        st.rerun()

# ── KICK OFF (after UI is rendered — avoids position issues) ──────────────────
if run_btn and topic.strip() and not st.session_state.pipeline_running:
    paper_paths = []
    if uploaded:
        for f in uploaded:
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
            tmp.write(f.read()); tmp.close()
            paper_paths.append(tmp.name)

    _prog.clear()
    import agents as _ag  # noqa: F401 — ensure module loaded

    state = initial_state(topic.strip(), paper_paths, user_id=_current_user_email)
    future = _fresh_executor().submit(_run, state)

    st.session_state.pipeline_future      = future
    st.session_state.pipeline_paper_paths = paper_paths
    st.session_state.pipeline_topic       = topic.strip()
    st.session_state.pipeline_log         = []
    st.session_state.pipeline_running     = True
    st.session_state.pop("pipeline_result", None)
    st.session_state.pop("pipeline_error", None)
    st.rerun()
